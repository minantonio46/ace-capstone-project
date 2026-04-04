from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── 페이지 설정 ──────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ─── 색상 상수 ────────────────────────────────────────────
COLOR_PRIMARY   = RGBColor(0x1F, 0x5C, 0x99)   # 진한 파랑
COLOR_SECONDARY = RGBColor(0x2E, 0x86, 0xAB)   # 중간 파랑
COLOR_ACCENT    = RGBColor(0xE8, 0xF4, 0xFD)   # 연한 파랑 배경
COLOR_GREEN     = RGBColor(0x27, 0x6E, 0x48)   # GET
COLOR_ORANGE    = RGBColor(0xB5, 0x5A, 0x00)   # POST
COLOR_RED       = RGBColor(0x8B, 0x00, 0x00)   # DELETE
COLOR_PURPLE    = RGBColor(0x5B, 0x2C, 0x6F)   # PUT
COLOR_GRAY_BG   = RGBColor(0xF5, 0xF5, 0xF5)
COLOR_DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

METHOD_COLORS = {
    "GET":    (RGBColor(0x27, 0x6E, 0x48), RGBColor(0xE8, 0xF5, 0xEE)),
    "POST":   (RGBColor(0x8B, 0x4A, 0x00), RGBColor(0xFF, 0xF0, 0xE0)),
    "PUT":    (RGBColor(0x5B, 0x2C, 0x6F), RGBColor(0xF5, 0xEA, 0xFF)),
    "DELETE": (RGBColor(0x7B, 0x00, 0x00), RGBColor(0xFF, 0xE8, 0xE8)),
}

# ─── 헬퍼: 셀 배경색 ──────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def rgb_to_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg_rgb(cell, rgb: RGBColor):
    set_cell_bg(cell, rgb_to_hex(rgb))

# ─── 헬퍼: 셀 테두리 ──────────────────────────────────────
def set_cell_border(cell, sides=("top","bottom","left","right"), color="CCCCCC", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

# ─── 헬퍼: 단락 스타일 ────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.bold  = True
    run.font.size  = Pt(16)
    run.font.color.rgb = COLOR_PRIMARY
    # 하단 테두리
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1F5C99")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.bold  = True
    run.font.size  = Pt(13)
    run.font.color.rgb = COLOR_SECONDARY
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.bold  = True
    run.font.size  = Pt(11)
    run.font.color.rgb = COLOR_DARK_GRAY
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_DARK_GRAY
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run("※  " + text)
    run.font.size  = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x88)
    return p

# ─── 헬퍼: HTTP 메서드 뱃지 ───────────────────────────────
def endpoint_header(method: str, path: str, summary: str):
    fg, bg = METHOD_COLORS.get(method, (COLOR_DARK_GRAY, COLOR_GRAY_BG))
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # 컬럼 너비
    widths = [Cm(1.8), Cm(8.0), Cm(6.2)]
    for i, w in enumerate(widths):
        tbl.columns[i].width = w
    row = tbl.rows[0]
    # 메서드 셀
    c0 = row.cells[0]
    set_cell_bg_rgb(c0, fg)
    set_cell_border(c0, color=rgb_to_hex(fg))
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(method)
    r0.font.bold  = True
    r0.font.size  = Pt(10)
    r0.font.color.rgb = COLOR_WHITE
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 경로 셀
    c1 = row.cells[1]
    set_cell_bg_rgb(c1, bg)
    set_cell_border(c1, color="DDDDDD")
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(path)
    r1.font.bold  = True
    r1.font.size  = Pt(10)
    r1.font.name  = "Consolas"
    r1.font.color.rgb = fg
    # 설명 셀
    c2 = row.cells[2]
    set_cell_bg(c2, "F8F8F8")
    set_cell_border(c2, color="DDDDDD")
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(summary)
    r2.font.size  = Pt(10)
    r2.font.color.rgb = COLOR_DARK_GRAY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ─── 헬퍼: 일반 표 (헤더+행) ──────────────────────────────
def data_table(headers: list, rows: list, col_widths_cm: list = None):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.style = "Table Grid"
    total_w = Cm(16)
    if col_widths_cm:
        widths = [Cm(w) for w in col_widths_cm]
    else:
        widths = [total_w // n] * n
    # 헤더 행
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.width = widths[i]
        set_cell_bg(c, "1F5C99")
        set_cell_border(c, color="1F5C99")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold  = True
        r.font.size  = Pt(9.5)
        r.font.color.rgb = COLOR_WHITE
    # 데이터 행
    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else "F5F8FC"
        for ci, val in enumerate(row_data):
            c = drow.cells[ci]
            c.width = widths[ci]
            set_cell_bg(c, bg)
            set_cell_border(c, color="CCCCCC")
            p = c.paragraphs[0]
            # 코드 스타일 (첫 번째 컬럼 등)
            if ci == 0:
                r = p.add_run(val)
                r.font.name  = "Consolas"
                r.font.size  = Pt(9)
                r.font.color.rgb = COLOR_SECONDARY
            else:
                r = p.add_run(val)
                r.font.size  = Pt(9.5)
                r.font.color.rgb = COLOR_DARK_GRAY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── 헬퍼: JSON 코드 블록 ─────────────────────────────────
def code_block(title: str, json_text: str):
    # 제목 단락
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after  = Pt(1)
    tr = tp.add_run(title)
    tr.font.bold  = True
    tr.font.size  = Pt(9.5)
    tr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    # 코드 박스
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    c = tbl.rows[0].cells[0]
    c.width = Cm(16)
    set_cell_bg(c, "1E1E1E")
    set_cell_border(c, color="333333")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    for line in json_text.strip().split("\n"):
        if p.runs:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name  = "Consolas"
        r.font.size  = Pt(9)
        r.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ──────────────────────────────────────────────────────────
#  표지
# ──────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Cm(4)
cover.paragraph_format.space_after  = Pt(4)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cover.add_run("스켈레톤 기반 의료 서비스")
cr.font.bold  = True
cr.font.size  = Pt(22)
cr.font.color.rgb = COLOR_PRIMARY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(8)
sr = sub.add_run("백엔드 API 명세서  v1.0")
sr.font.size  = Pt(14)
sr.font.color.rgb = COLOR_SECONDARY

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run("2026-04-01")
dr.font.size  = Pt(11)
dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ──────────────────────────────────────────────────────────
# 1. 개요
# ──────────────────────────────────────────────────────────
h1("1. 개요")

body("이 문서는 스켈레톤 기반 의료 서비스 백엔드(Spring Boot 3.2.3)가 제공하는 REST API와 WebSocket 인터페이스를 기술합니다.")
body("프론트엔드 개발팀이 참고하여 간호사·의사용 앱을 구현할 수 있도록 작성되었습니다.")

h2("1.1 기술 스택")
data_table(
    ["구분", "기술"],
    [
        ["프레임워크", "Spring Boot 3.2.3 / Java 17"],
        ["인증",       "Spring Security + JWT (Bearer Token)"],
        ["실시간 알림", "Spring WebSocket + STOMP"],
        ["데이터베이스", "MySQL 8 + JPA/Hibernate"],
        ["API 문서",   "Swagger UI  →  http://localhost:8080/swagger-ui.html"],
    ],
    col_widths_cm=[4, 12]
)

h2("1.2 아키텍처")
body("CCTV → AI 서버(COME-HVLM) → Spring Boot 백엔드 → 프론트엔드 앱")
body("AI 서버는 스켈레톤 분석 결과만 백엔드에 전송하며, 영상 원본은 백엔드를 거치지 않습니다.")

# ──────────────────────────────────────────────────────────
# 2. 공통 규약
# ──────────────────────────────────────────────────────────
h1("2. 공통 규약")

h2("2.1 Base URL")
body("http://localhost:8080   (개발 환경 기준)")

h2("2.2 인증")
body("로그인 후 발급된 Access Token을 모든 요청 헤더에 포함해야 합니다.")
body("Authorization: Bearer {accessToken}", indent=1)
note("토큰 만료 시간: 1시간 / Refresh Token 유효 기간: 7일")

h2("2.3 공통 응답 형식")
code_block("성공 응답", """{
  "success": true,
  "message": null,
  "data": { ... }
}""")
code_block("실패 응답", """{
  "success": false,
  "message": "오류 설명",
  "data": null
}""")

h2("2.4 HTTP 상태 코드")
data_table(
    ["코드", "의미", "설명"],
    [
        ["200 OK",         "성공",           "조회·수정·삭제 성공"],
        ["201 Created",    "생성 성공",       "등록 성공"],
        ["400 Bad Request","잘못된 요청",     "파라미터 오류, 유효성 검증 실패"],
        ["401 Unauthorized","인증 실패",      "토큰 없음 또는 만료"],
        ["403 Forbidden",  "권한 없음",       "역할(Role) 미충족"],
        ["404 Not Found",  "리소스 없음",     "존재하지 않는 ID 등"],
        ["500 Internal Server Error","서버 오류", "예상치 못한 서버 내부 오류"],
    ],
    col_widths_cm=[3.5, 3.5, 9.0]
)

h2("2.5 날짜/시간 형식")
body("ISO 8601 형식을 사용합니다:  yyyy-MM-dd'T'HH:mm:ss  (서버 타임존: Asia/Seoul)")
body("예)  2026-04-01T09:30:00", indent=1)

h2("2.6 역할(Role)")
data_table(
    ["역할", "설명"],
    [
        ["NURSE", "간호사 — 바이탈 기록, 약물 투여, 메모 작성"],
        ["DOCTOR", "의사 — 처방, 환자 등록"],
        ["ADMIN", "관리자 — 전체 권한"],
    ],
    col_widths_cm=[4, 12]
)

# ──────────────────────────────────────────────────────────
# 3. 인증 API
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("3. 인증 API")
note("인증 API는 토큰 없이 호출합니다.")

# 3.1 로그인
h2("3.1 로그인")
endpoint_header("POST", "/api/auth/login", "아이디/비밀번호로 JWT 토큰 발급")

h3("Request Body")
data_table(
    ["필드", "타입", "필수", "설명"],
    [
        ["username", "String", "Y", "사용자 아이디"],
        ["password", "String", "Y", "비밀번호"],
    ],
    col_widths_cm=[3.5, 2.5, 1.5, 8.5]
)
code_block("Request 예시", """{
  "username": "nurse01",
  "password": "password123"
}""")
code_block("Response 예시", """{
  "success": true,
  "data": {
    "accessToken":  "eyJhbGci...",
    "refreshToken": "eyJhbGci...",
    "username": "nurse01",
    "name":     "김간호",
    "role":     "NURSE"
  }
}""")

# 3.2 토큰 갱신
h2("3.2 토큰 갱신")
endpoint_header("POST", "/api/auth/refresh", "Refresh Token으로 새 Access Token 발급")

h3("Request Header")
data_table(
    ["헤더", "타입", "필수", "설명"],
    [
        ["Refresh-Token", "String", "Y", "로그인 시 발급된 Refresh Token"],
    ],
    col_widths_cm=[4, 2, 1.5, 8.5]
)
note("Response 구조는 로그인과 동일합니다.")

# 3.3 로그아웃
h2("3.3 로그아웃")
endpoint_header("POST", "/api/auth/logout", "로그아웃 (클라이언트 측 토큰 삭제)")
note("서버는 stateless이므로, 클라이언트에서 저장된 토큰을 삭제하면 됩니다. 응답: 200 OK")

# ──────────────────────────────────────────────────────────
# 4. 환자 관리 API
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("4. 환자 관리 API")

h2("4.1 병동 환자 목록 조회")
endpoint_header("GET", "/api/patients?wardId={wardId}", "입원 중인 환자 전체 목록")

h3("Query Parameter")
data_table(
    ["파라미터", "타입", "필수", "설명"],
    [["wardId", "Long", "Y", "병동 ID"]],
    col_widths_cm=[3.5, 2.5, 1.5, 8.5]
)
code_block("Response 예시", """{
  "success": true,
  "data": [
    {
      "id": 1,
      "name": "홍길동",
      "birthDate": "1965-03-15",
      "gender": "M",
      "wardId": 1,
      "bedNumber": "302-1",
      "admissionDate": "2026-03-20",
      "dischargeDate": null,
      "status": "ADMITTED",
      "doctorId": 5,
      "createdAt": "2026-03-20T09:00:00"
    }
  ]
}""")

h2("4.2 환자 상세 조회")
endpoint_header("GET", "/api/patients/{id}", "특정 환자 상세 정보")

h2("4.3 환자 등록")
endpoint_header("POST", "/api/patients", "신규 환자 등록")
h3("Request Body")
data_table(
    ["필드", "타입", "필수", "설명"],
    [
        ["name",          "String",    "Y", "환자 이름"],
        ["birthDate",     "LocalDate", "Y", "생년월일 (yyyy-MM-dd)"],
        ["gender",        "String",    "Y", "성별 (M / F)"],
        ["wardId",        "Long",      "Y", "병동 ID"],
        ["bedNumber",     "String",    "Y", "병상 번호 (예: 302-1)"],
        ["admissionDate", "LocalDate", "Y", "입원일 (yyyy-MM-dd)"],
        ["doctorId",      "Long",      "N", "담당 의사 ID"],
    ],
    col_widths_cm=[3.5, 2.5, 1.5, 8.5]
)

h2("4.4 환자 정보 수정")
endpoint_header("PUT", "/api/patients/{id}", "이름, 병상, 담당의 수정")

h2("4.5 퇴원 처리")
endpoint_header("DELETE", "/api/patients/{id}", "퇴원 처리 (status → DISCHARGED)")
note("실제 삭제가 아니라 status를 DISCHARGED로 변경합니다.")

h3("환자 Status 값")
data_table(
    ["값", "의미"],
    [
        ["ADMITTED",   "입원중"],
        ["DISCHARGED", "퇴원"],
        ["IN_SURGERY", "수술중"],
    ],
    col_widths_cm=[4, 12]
)

# ──────────────────────────────────────────────────────────
# 5. 약물 관리 API
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("5. 약물 관리 API")
note("투여 기한 초과 시 WebSocket으로 자동 알림이 발송됩니다.")

h2("5.1 처방 약물 목록 조회")
endpoint_header("GET", "/api/patients/{patientId}/medications", "현재 활성 처방 약물 목록")

code_block("Response 예시", """{
  "success": true,
  "data": [
    {
      "id": 10,
      "patientId": 1,
      "drugName": "혈압약 (암로디핀)",
      "dosage": 5.0,
      "unit": "mg",
      "intervalHours": 24,
      "startAt": "2026-03-20T09:00:00",
      "endAt": null,
      "prescribedBy": 5,
      "status": "ACTIVE",
      "createdAt": "2026-03-20T09:00:00"
    }
  ]
}""")

h2("5.2 약물 처방 추가")
endpoint_header("POST", "/api/patients/{patientId}/medications", "새 약물 처방 등록 (의사 권한)")
h3("Request Body")
data_table(
    ["필드", "타입", "필수", "설명"],
    [
        ["drugName",      "String",  "Y", "약물명"],
        ["dosage",        "Double",  "Y", "용량 (숫자)"],
        ["unit",          "String",  "N", "단위 (mg, ml 등)"],
        ["intervalHours", "Integer", "Y", "투여 주기 (시간 단위, 예: 8)"],
        ["startAt",       "LocalDateTime", "Y", "처방 시작일시"],
        ["endAt",         "LocalDateTime", "N", "처방 종료일시 (없으면 무기한)"],
        ["prescribedBy",  "Long",    "Y", "처방 의사 ID"],
    ],
    col_widths_cm=[3.5, 2.5, 1.5, 8.5]
)

h2("5.3 약물 투여 기록")
endpoint_header("POST", "/api/patients/{patientId}/medications/{medicationId}/administer", "투여 완료 기록")
h3("Request Body")
data_table(
    ["필드", "타입", "필수", "설명"],
    [
        ["administeredBy", "Long",   "Y", "투여한 간호사 ID"],
        ["note",           "String", "N", "특이사항 메모"],
    ],
    col_widths_cm=[3.5, 2.5, 1.5, 8.5]
)
code_block("Response 예시 (MedicationRecord)", """{
  "success": true,
  "data": {
    "id": 55,
    "medicationId": 10,
    "patientId": 1,
    "drugName": "혈압약 (암로디핀)",
    "dosage": 5.0,
    "administeredAt": "2026-04-01T09:30:00",
    "nextDueAt": "2026-04-02T09:30:00",
    "administeredBy": 3,
    "note": null
  }
}""")
note("동일 약물을 주기의 절반보다 짧은 시간 내에 재투여하려 하면 400 오류가 반환됩니다.")

h2("5.4 오늘 투여 스케줄")
endpoint_header("GET", "/api/patients/{patientId}/medications/schedule", "오늘 기준 투여 예정 목록")

# ──────────────────────────────────────────────────────────
# 6. 바이탈 관리 API
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("6. 바이탈 관리 API")
note("비정상 수치 감지 시 WebSocket으로 즉시 경고 알림이 발송됩니다.")

h2("6.1 바이탈 기록")
endpoint_header("POST", "/api/patients/{patientId}/vitals", "체온·혈압 등 바이탈 기록")
h3("Request Body")
data_table(
    ["필드", "타입", "필수", "설명"],
    [
        ["temperature",      "Double",  "N", "체온 (°C)"],
        ["bpSystolic",       "Integer", "N", "혈압 수축기 (mmHg)"],
        ["bpDiastolic",      "Integer", "N", "혈압 이완기 (mmHg)"],
        ["heartRate",        "Integer", "N", "맥박 (bpm)"],
        ["oxygenSaturation", "Integer", "N", "산소포화도 (%)"],
        ["respiratoryRate",  "Integer", "N", "호흡수 (회/분)"],
        ["recordedBy",       "Long",    "Y", "기록한 간호사 ID"],
        ["note",             "String",  "N", "특이사항"],
    ],
    col_widths_cm=[3.8, 2.2, 1.5, 8.5]
)

h3("이상 감지 기준")
data_table(
    ["항목", "이상 범위", "알림 등급"],
    [
        ["체온",       "> 38.0°C",           "HIGH"],
        ["혈압 수축기", "> 140 또는 < 90 mmHg", "HIGH"],
        ["산소포화도",  "< 95%",              "HIGH"],
        ["맥박",       "> 100 또는 < 60 bpm", "HIGH"],
    ],
    col_widths_cm=[3.5, 5, 7.5]
)

h2("6.2 최신 바이탈 조회")
endpoint_header("GET", "/api/patients/{patientId}/vitals/latest", "가장 최근 바이탈 1건")

h2("6.3 기간별 바이탈 이력")
endpoint_header("GET", "/api/patients/{patientId}/vitals/history", "기간 내 바이탈 목록 (시계열)")
h3("Query Parameters")
data_table(
    ["파라미터", "타입", "필수", "예시"],
    [
        ["from", "LocalDateTime", "Y", "2026-04-01T00:00:00"],
        ["to",   "LocalDateTime", "Y", "2026-04-01T23:59:59"],
    ],
    col_widths_cm=[3.5, 3, 1.5, 8]
)

# ──────────────────────────────────────────────────────────
# 7. 메모 관리 API
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("7. 메모 관리 API")

h2("7.1 메모 목록 조회")
endpoint_header("GET", "/api/patients/{patientId}/notes", "환자 메모 전체 목록 (최신순)")

h2("7.2 메모 작성")
endpoint_header("POST", "/api/patients/{patientId}/notes", "새 메모 등록")
h3("Request Body")
data_table(
    ["필드", "타입", "필수", "설명"],
    [
        ["noteType",  "String", "Y", "PATIENT / GUARDIAN / CAUTION 중 1개"],
        ["content",   "String", "Y", "메모 내용"],
        ["createdBy", "Long",   "Y", "작성자 사용자 ID"],
    ],
    col_widths_cm=[3.5, 2.5, 1.5, 8.5]
)

h3("noteType 값")
data_table(
    ["값", "의미"],
    [
        ["PATIENT",  "환자 관련 메모"],
        ["GUARDIAN", "보호자 관련 메모"],
        ["CAUTION",  "주의사항 (알레르기, 특이사항 등)"],
    ],
    col_widths_cm=[4, 12]
)

h2("7.3 메모 수정")
endpoint_header("PUT", "/api/patients/{patientId}/notes/{noteId}", "메모 내용 수정")

h2("7.4 메모 삭제")
endpoint_header("DELETE", "/api/patients/{patientId}/notes/{noteId}", "메모 삭제")

# ──────────────────────────────────────────────────────────
# 8. 병동 대시보드 API
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("8. 병동 대시보드 API")

h2("8.1 병동 전체 현황 조회")
endpoint_header("GET", "/api/wards/{wardId}/dashboard", "간호사 메인 화면용 병동 전체 요약")

code_block("Response 예시", """{
  "success": true,
  "data": {
    "wardId": 1,
    "wardName": "3층 일반병동",
    "totalPatients": 12,
    "patients": [
      {
        "patientId": 1,
        "name": "홍길동",
        "bedNumber": "302-1",
        "latestVitals": {
          "bp": "130/85",
          "temperature": 37.2,
          "oxygenSaturation": 98,
          "heartRate": 72,
          "recordedAt": "2026-04-01T08:00:00"
        },
        "nextMedication": {
          "drugName": "혈압약 (암로디핀)",
          "minutesLeft": 45,
          "nextDueAt": "2026-04-01T10:15:00"
        },
        "skeletonStatus": "NORMAL",
        "unresolvedAlerts": 0
      }
    ]
  }
}""")

h3("skeletonStatus 값")
data_table(
    ["값", "의미"],
    [
        ["NORMAL",       "정상 활동 감지"],
        ["NO_MOVEMENT",  "움직임 없음"],
        ["FALL_DETECTED","낙상 감지 (긴급)"],
    ],
    col_widths_cm=[4, 12]
)
note("latestVitals 또는 nextMedication이 없으면 해당 필드가 null로 반환됩니다.")

# ──────────────────────────────────────────────────────────
# 9. 알림 관리 API
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("9. 알림 관리 API")
note("알림은 WebSocket으로 실시간 수신되며, 이 API는 이력 조회 및 해결 처리에 사용합니다.")

h2("9.1 병동 미해결 알림 목록")
endpoint_header("GET", "/api/alerts/ward/{wardId}", "특정 병동의 미해결 알림 (최신순)")

h2("9.2 환자 미해결 알림 목록")
endpoint_header("GET", "/api/alerts/patient/{patientId}", "특정 환자의 미해결 알림")

h2("9.3 알림 해결 처리")
endpoint_header("POST", "/api/alerts/{alertId}/resolve?userId={userId}", "알림을 해결 완료 처리")

h3("Alert 객체 구조")
data_table(
    ["필드", "타입", "설명"],
    [
        ["id",          "Long",          "알림 ID"],
        ["patientId",   "Long",          "환자 ID"],
        ["wardId",      "Long",          "병동 ID"],
        ["alertType",   "String",        "알림 유형 (아래 표 참고)"],
        ["severity",    "String",        "심각도 (INFO / WARNING / HIGH / CRITICAL)"],
        ["message",     "String",        "알림 메시지"],
        ["resolved",    "Boolean",       "해결 여부"],
        ["resolvedBy",  "Long",          "해결한 사용자 ID"],
        ["resolvedAt",  "LocalDateTime", "해결 일시"],
        ["occurredAt",  "LocalDateTime", "발생 일시"],
    ],
    col_widths_cm=[3.5, 3, 9.5]
)

h3("alertType 값")
data_table(
    ["값", "트리거", "심각도"],
    [
        ["FALL_DETECTED",       "낙상 감지 (AI 서버)",    "CRITICAL"],
        ["NO_MOVEMENT",         "장시간 움직임 없음 (30분+)", "WARNING"],
        ["MEDICATION_OVERDUE",  "약물 투여 기한 초과",    "WARNING"],
        ["MEDICATION_DUE_SOON", "투여 30분 전 알림",       "INFO"],
        ["ABNORMAL_VITAL",      "비정상 바이탈 감지",      "HIGH"],
    ],
    col_widths_cm=[4.5, 6, 5.5]
)

# ──────────────────────────────────────────────────────────
# 10. 카메라 관리 API
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("10. 카메라 관리 API")

h2("10.1 카메라 목록 조회")
endpoint_header("GET", "/api/cameras", "전체 CCTV 카메라 목록")

h2("10.2 환자 배정")
endpoint_header("POST", "/api/cameras/{id}/assign?patientId={patientId}", "카메라에 환자 배정")
note("한 카메라에는 한 명의 환자만 배정됩니다. 기존 배정 환자가 있으면 덮어씁니다.")

h2("10.3 환자 배정 해제")
endpoint_header("POST", "/api/cameras/{id}/unassign", "카메라 배정 해제")

h3("Camera 객체 구조")
data_table(
    ["필드", "타입", "설명"],
    [
        ["id",                "Long",   "카메라 ID"],
        ["cameraCode",        "String", "카메라 코드 (예: CAM_101)"],
        ["wardId",            "Long",   "소속 병동 ID"],
        ["location",          "String", "설치 위치 (예: 302호)"],
        ["status",            "String", "ACTIVE / INACTIVE / ERROR"],
        ["assignedPatientId", "Long",   "배정된 환자 ID (없으면 null)"],
    ],
    col_widths_cm=[3.8, 2.5, 9.7]
)

# ──────────────────────────────────────────────────────────
# 11. AI 서버 이벤트 API
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("11. AI 서버 이벤트 API")
note("이 API는 AI 서버에서 백엔드로 호출합니다. 프론트엔드에서는 호출하지 않습니다.")

h2("11.1 스켈레톤 분석 이벤트 수신")
endpoint_header("POST", "/api/ai/skeleton-event", "AI 서버가 분석 결과를 전송하는 엔드포인트")

h3("인증")
body("JWT 토큰 대신 내부망 전용 API Key를 사용합니다.")
body("X-AI-API-Key: {apiKey}  (application.yml의 ai-server.api-key 값)", indent=1)

h3("Request Body")
data_table(
    ["필드", "타입", "필수", "설명"],
    [
        ["cameraId",   "String",        "Y", "카메라 코드 (예: CAM_101)"],
        ["patientId",  "Long",          "N", "환자 ID (없으면 카메라 배정에서 자동 조회)"],
        ["eventType",  "String",        "Y", "FALL_DETECTED / NO_MOVEMENT / NORMAL"],
        ["timestamp",  "LocalDateTime", "Y", "이벤트 발생 시각"],
        ["confidence", "Double",        "Y", "신뢰도 (0.0 ~ 1.0)"],
    ],
    col_widths_cm=[3.5, 3, 1.5, 8]
)
code_block("Request 예시", """{
  "cameraId":   "CAM_101",
  "eventType":  "FALL_DETECTED",
  "timestamp":  "2026-04-01T14:30:00",
  "confidence": 0.95
}""")

# ──────────────────────────────────────────────────────────
# 12. WebSocket 명세
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("12. WebSocket 실시간 알림 명세")

h2("12.1 연결 엔드포인트")
body("ws://localhost:8080/ws  (SockJS fallback 지원)")

h2("12.2 연결 예시 (JavaScript)")
code_block("SockJS + STOMP 연결 코드", """import SockJS from 'sockjs-client';
import { Client } from '@stomp/stompjs';

const client = new Client({
  webSocketFactory: () => new SockJS('http://localhost:8080/ws'),
  onConnect: () => {
    // 병동 ID 1번 알림 구독
    client.subscribe('/topic/ward/1/alerts', (msg) => {
      const alert = JSON.parse(msg.body);
      console.log('알림 수신:', alert);
    });
  },
});
client.activate();""")

h2("12.3 구독 경로")
data_table(
    ["경로", "설명"],
    [
        ["/topic/ward/{wardId}/alerts", "해당 병동의 모든 실시간 알림 수신"],
    ],
    col_widths_cm=[6, 10]
)

h2("12.4 알림 메시지 형식")
code_block("수신 메시지 예시 (AlertMessage)", """{
  "type":        "FALL_DETECTED",
  "patientId":   1,
  "patientName": "홍길동",
  "bedNumber":   "302-1",
  "message":     "낙상이 감지되었습니다! 즉시 확인이 필요합니다.",
  "severity":    "CRITICAL",
  "occurredAt":  "2026-04-01T14:30:00"
}""")

h3("심각도(severity) 별 처리 가이드")
data_table(
    ["severity", "색상 권장", "대응"],
    [
        ["CRITICAL", "빨강 (#D32F2F)", "즉시 팝업 + 음성 알림"],
        ["HIGH",     "주황 (#F57C00)", "상단 배너 표시"],
        ["WARNING",  "노랑 (#FBC02D)", "알림 목록에 추가"],
        ["INFO",     "파랑 (#1976D2)", "토스트 메시지"],
    ],
    col_widths_cm=[3.5, 4, 8.5]
)

# ──────────────────────────────────────────────────────────
# 13. 에러 처리 가이드
# ──────────────────────────────────────────────────────────
doc.add_page_break()
h1("13. 에러 처리 가이드")

h2("13.1 공통 에러 응답 처리")
code_block("에러 응답 예시", """{
  "success": false,
  "message": "파라미터 오류: name: 이름은 필수입니다",
  "data": null
}""")

h2("13.2 주요 에러 시나리오")
data_table(
    ["시나리오", "HTTP 상태", "message 예시"],
    [
        ["로그인 실패",          "401", "아이디 또는 비밀번호가 올바르지 않습니다."],
        ["토큰 만료",            "401", "(빈 응답 — 필터에서 차단)"],
        ["권한 없음",            "403", "접근 권한이 없습니다."],
        ["존재하지 않는 환자",   "400", "환자를 찾을 수 없습니다: 999"],
        ["중복 투여 시도",       "400", "이전 투여 후 너무 짧은 시간입니다. 다음 최소 투여 가능 시간: ..."],
        ["필수 필드 누락",       "400", "name: 이름은 필수입니다, wardId: ..."],
        ["존재하지 않는 카메라", "400", "등록되지 않은 카메라: CAM_999"],
        ["AI API Key 불일치",   "403", "인증 실패"],
    ],
    col_widths_cm=[4.5, 2.5, 9]
)

h2("13.3 토큰 갱신 플로우")
body("1.  API 요청 → 401 응답 수신")
body("2.  저장된 Refresh Token으로  POST /api/auth/refresh  호출")
body("3.  새 Access Token 저장 후 원래 요청 재시도")
body("4.  Refresh Token도 만료됐다면 → 로그인 화면으로 이동")

# ──────────────────────────────────────────────────────────
# 저장
# ──────────────────────────────────────────────────────────
out_path = "C:/스켈레톤기반의료서비스/API_명세서_v1.0.docx"
doc.save(out_path)
print(f"생성 완료: {out_path}")
